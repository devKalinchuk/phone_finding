#!/usr/bin/env python3
import argparse
import re
from pathlib import Path
from docx import Document


TARGET_FOLDER = "."  # Папка за замовчуванням

# ANSI-коди для підсвічування тексту червоним кольором у терміналі
RED = "\033[91m"
RESET = "\033[0m"


def extract_paragraphs(doc):
    """
    Генератор, який повертає всі абзаци документа (Paragraph-об'єкти):
    основний текст, текст у таблицях та колонтитулах.
    """
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

    for section in doc.sections:
        if section.header:
            yield from section.header.paragraphs
        if section.footer:
            yield from section.footer.paragraphs


def term_to_regex(term):
    """
    Перетворює один пошуковий термін на фрагмент регулярного виразу.
    Спочатку весь термін екранується (re.escape), щоб спецсимволи
    сприймались буквально, а потім екрановані ? та * повертаються назад
    у вигляді підстановочних знаків:
      ?  -> будь-який один символ (крім пробілу)
      *  -> будь-яка кількість (0+) символів у межах слова (без пробілів)
    Підстановочні знаки обмежені межами слова, щоб "пошук*" не
    "з'їдав" сусідні слова в реченні.
    """
    escaped = re.escape(term)
    escaped = escaped.replace(r"\?", r"\S").replace(r"\*", r"\S*")
    return escaped


def build_patterns(terms, case_sensitive):
    """
    Створює:
    1) individual_patterns - список окремих скомпільованих виразів,
       по одному на кожен термін (слово чи фразу з лапок) - потрібні,
       щоб перевірити, що ВСІ терміни присутні в документі;
    2) combined_pattern - один вираз-«або» з усіх термінів - потрібен
       лише для підсвічування збігів у знайдених абзацах.
    Символи ? та * у будь-якому терміні працюють як підстановочні знаки.
    """
    flags = 0 if case_sensitive else re.IGNORECASE
    individual_patterns = [re.compile(term_to_regex(term), flags) for term in terms]
    combined_pattern = re.compile("|".join(term_to_regex(term) for term in terms), flags)
    return individual_patterns, combined_pattern


def highlight(text, pattern):
    """Підсвічує всі знайдені збіги червоним кольором для виводу в термінал."""
    return pattern.sub(lambda m: f"{RED}{m.group(0)}{RESET}", text)


def find_text_in_folder(folder_path_str, terms, case_sensitive=False, recursive=False):
    folder = Path(folder_path_str)

    if not folder.exists() or not folder.is_dir():
        print(f"Помилка: Шлях '{folder}' не існує або не є папкою.")
        return

    glob_method = folder.rglob if recursive else folder.glob
    docx_files = [
        f for f in glob_method("*.docx")
        if not f.name.startswith("~$")
    ]

    if not docx_files:
        print("У вказаній папці не знайдено документів .docx")
        return

    individual_patterns, combined_pattern = build_patterns(terms, case_sensitive)
    total_matches = 0

    for file_path in docx_files:
        try:
            doc = Document(file_path)

            # Спочатку зчитуємо всі непорожні абзаци один раз
            paragraph_texts = [
                para.text for para in extract_paragraphs(doc) if para.text.strip()
            ]
            full_text = "\n".join(paragraph_texts)

            # Документ підходить, лише якщо КОЖЕН термін пошуку
            # знайдений хоча б десь у документі
            if not all(p.search(full_text) for p in individual_patterns):
                continue

            file_header_printed = False
            for text in paragraph_texts:
                if combined_pattern.search(text):
                    if not file_header_printed:
                        print(f"📄 Файл: {file_path.resolve()}")
                        file_header_printed = True
                    print(f"   {highlight(text, combined_pattern)}")
                    total_matches += 1

            if file_header_printed:
                print("-" * 40)

        except Exception as e:
            print(f"⚠️ Не вдалося прочитати файл {file_path.name}: {e}")

    if total_matches == 0:
        print("Збігів не знайдено.")


def parse_args():
    parser = argparse.ArgumentParser(
        prog="finder",
        description="Пошук тексту в документах .docx з підсвічуванням збігів червоним кольором."
    )
    parser.add_argument(
        "query",
        nargs="+",
        help='Текст для пошуку. Якщо вказано кілька слів чи фраз (finder слово1 "фраза 2"), '
             'у документі мають бути присутні УСІ вони одночасно (не обов\'язково в одному абзаці). '
             'Текст "у лапках" шукається як одна нерозривна фраза. '
             'Підтримуються підстановочні знаки: ? - будь-який один символ, '
             '* - будь-яка кількість символів (наприклад: finder тел?фон, finder пошук*).'
    )
    parser.add_argument(
        "-c", "--case-sensitive",
        action="store_true",
        help="Точний пошук з урахуванням регістру (за замовчуванням регістр ігнорується)."
    )
    parser.add_argument(
        "-f", "--folder",
        default=TARGET_FOLDER,
        help="Папка з документами .docx (за замовчуванням поточна папка).",
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="Шукати документи також у підпапках.",
    )
    return parser.parse_args()


# --- Запуск ---
if __name__ == "__main__":
    args = parse_args()
    print(
        f"Пошук: {' | '.join(args.query)}\n"
        f"Папка: '{Path(args.folder).resolve()}'"
        f"{' (рекурсивно)' if args.recursive else ''}"
        f"{' [точний регістр]' if args.case_sensitive else ''}\n"
    )
    find_text_in_folder(
        args.folder,
        args.query,
        case_sensitive=args.case_sensitive,
        recursive=args.recursive,
    )