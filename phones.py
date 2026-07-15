import argparse
from pathlib import Path
import re
from docx import Document


# Регулярний вираз для пошуку українських номерів телефонів
PHONE_PATTERN = re.compile(
    r'(?:\+?38)?\s*\(?0\d{2}\)?[-.\s]*\d{3}[-.\s]*\d{2}[-.\s]*\d{2}\b'
)

# Регулярний вираз для пошуку електронних адрес
EMAIL_PATTERN = re.compile(
    r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}'
)


def extract_text_sources(doc):
    """
    Генератор, який послідовно повертає весь текст із різних частин документа:
    основні абзаци, таблиці та колонтитули.
    """
    # 1. Текст з основних абзаців
    for para in doc.paragraphs:
        yield para.text

    # 2. Текст із таблиць
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.text

    # 3. Текст із колонтитулів (якщо вони існують)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                yield para.text
        if section.footer:
            for para in section.footer.paragraphs:
                yield para.text


def find_contacts_in_folder(folder_path_str, recursive=False):
    folder = Path(folder_path_str)

    if not folder.exists() or not folder.is_dir():
        print(f"Помилка: Шлях '{folder}' не існує або не є папкою.")
        return

    # Обираємо метод пошуку файлів залежно від прапорця --recursive
    glob_method = folder.rglob if recursive else folder.glob

    docx_files = [
        f for f in glob_method("*.docx")
        if not f.name.startswith("~$")
    ]

    if not docx_files:
        print("У вказаній папці не знайдено документів .docx")
        return

    for file_path in docx_files:
        try:
            doc = Document(file_path)
            phones_in_file = []
            emails_in_file = []

            # Збираємо текст з усіх джерел у документі
            for text in extract_text_sources(doc):
                if not text.strip():
                    continue

                for match in PHONE_PATTERN.findall(text):
                    cleaned_phone = match.strip()
                    if cleaned_phone and cleaned_phone not in phones_in_file:
                        phones_in_file.append(cleaned_phone)

                for match in EMAIL_PATTERN.findall(text):
                    cleaned_email = match.strip()
                    if cleaned_email and cleaned_email not in emails_in_file:
                        emails_in_file.append(cleaned_email)

            # Вивід результатів (лише якщо щось знайдено)
            if phones_in_file or emails_in_file:
                print(f"📄 Файл: {file_path.name}")
                for phone in phones_in_file:
                    print(f"   📞 {phone}")
                for email in emails_in_file:
                    print(f"   ✉️ {email}")
                print("-" * 40)

        except Exception as e:
            print(f"⚠️ Не вдалося прочитати файл {file_path.name}: {e}")


def parse_args():
    parser = argparse.ArgumentParser(
        description="Пошук телефонних номерів та електронних адрес у документах .docx"
    )
    parser.add_argument(
        "folder",
        nargs="?",
        default=".",
        help="Шлях до папки з документами .docx (за замовчуванням поточна папка)",
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="Шукати документи також у підпапках",
    )
    return parser.parse_args()


# --- Запуск ---
if __name__ == "__main__":
    args = parse_args()
    print(f"Пошук у папці '{Path(args.folder).resolve()}'"
          f"{' (рекурсивно)' if args.recursive else ''}...\n")
    find_contacts_in_folder(args.folder, recursive=args.recursive)